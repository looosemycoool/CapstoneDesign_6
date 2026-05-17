import os
import sys
import json
import zipfile
import subprocess
import shutil
import platform
import tempfile
import opendataloader_pdf
import pandas as pd
from docx import Document

# Windows 콘솔(cp949)에서도 한글/유니코드 출력 안전하게
try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass

# 절대 경로로 설정
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
ATTACHMENT_DIR = os.path.join(BASE_DIR, "data", "attachments")
PARSED_DIR = os.path.join(BASE_DIR, "data", "parsed")
RAW_DIR = os.path.join(BASE_DIR, "data", "raw")
TMP_DIR = os.path.join(BASE_DIR, "data", "tmp_convert")

os.makedirs(PARSED_DIR, exist_ok=True)
os.makedirs(TMP_DIR, exist_ok=True)

# 추출 텍스트가 이 값보다 적으면 스캔본/변환 실패 가능성 경고
EMPTY_TEXT_THRESHOLD = 50


def find_libreoffice():
    """LibreOffice 실행 파일 탐색.
    1) PATH의 soffice 2) Windows 기본 설치 경로 3) None
    """
    found = shutil.which("soffice")
    if found:
        return found
    if platform.system() == "Windows":
        candidates = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ]
        for c in candidates:
            if os.path.exists(c):
                return c
    return None


LIBREOFFICE = find_libreoffice()


# ── 트랙별 졸업요건 표 어노테이션 ─────────────────────────────
# opendataloader가 PDF 표를 선형화하면 다중전공트랙/해외복수학위트랙/학석사연계트랙의
# 요건이 한 텍스트에 뒤섞인다. 각 ▸ 항목 앞에 [트랙명]을 삽입해
# LLM이 어느 트랙의 요건인지 구분할 수 있도록 한다.

_GRAD_TRACKS = [
    "다중전공트랙",
    "해외복수학위트랙",
    "학석사연계트랙",
    "학·석사연계트랙",
    "학・석사연계트랙",
]
_TRACK_CANONICAL = {
    "학·석사연계트랙":  "학석사연계트랙",
    "학・석사연계트랙": "학석사연계트랙",
}


def _annotate_track_sections(text: str) -> str:
    """PDF 표 선형화로 혼재된 트랙별 졸업요건에 '[트랙명]' 접두어 삽입.

    알고리즘:
      1. 트랙명 등장 줄 위치를 수집한다.
      2. 줄 → 트랙 매핑: 트랙명 이후는 forward fill,
         첫 트랙 등장 이전 줄은 첫 트랙으로 소급 적용한다.
      3. '▸'가 포함된 줄에 트랙 레이블이 없으면 '[트랙명]'을 앞에 붙인다.

    트랙명이 없는 문서는 원문 그대로 반환(no-op).
    """
    if not any(t in text for t in _GRAD_TRACKS):
        return text

    def canonical(name: str) -> str:
        return _TRACK_CANONICAL.get(name, name)

    lines = text.split('\n')

    # 1단계: (줄 번호, 정규 트랙명) 수집
    track_at: list[tuple[int, str]] = []
    for i, line in enumerate(lines):
        for t in _GRAD_TRACKS:
            if t in line:
                track_at.append((i, canonical(t)))
                break

    if not track_at:
        return text

    # 2단계: 줄 → 담당 트랙 매핑
    #   첫 트랙 등장 이전 줄은 첫 트랙으로 소급 (표 앞부분 요건도 첫 트랙 소속)
    first_line, first_track = track_at[0]
    line_track: dict[int, str] = {}
    for i in range(first_line):
        line_track[i] = first_track
    current, ptr = first_track, 0
    for i in range(first_line, len(lines)):
        if ptr < len(track_at) and i == track_at[ptr][0]:
            current = track_at[ptr][1]
            ptr += 1
        line_track[i] = current

    # 3단계: ▸ 항목에 레이블 삽입
    result = []
    for i, line in enumerate(lines):
        stripped = line.strip()
        if '▸' in stripped:
            track = line_track.get(i)
            if track and track not in stripped:
                line = f"[{track}] {stripped}"
        result.append(line)

    return '\n'.join(result)


def _annotate_markdown_table_rows(text: str) -> str:
    """Markdown pipe 표의 각 데이터 행에 컬럼 헤더를 붙여 자기완결적으로 만든다.

    청킹 시 헤더 행이 분리되어도 각 데이터 행이 의미를 보존하도록
    '[행레이블] 헤더1: 값1 | 헤더2: 값2 ...' 형식으로 변환한다.
    첫 번째 컬럼이 범주형(비수치) 값이면 [레이블]로 앞에 붙이고
    나머지는 헤더명: 값 쌍으로 나열한다.
    구분자 행(|---|)을 기준으로 표를 감지하므로 오탐 없음.
    헤더 행·구분자 행 원본은 유지해 keyword 검색에도 걸리도록 한다.
    """
    def parse_cells(line: str) -> list:
        return [c.strip() for c in line.strip().strip('|').split('|')]

    def is_separator_row(line: str) -> bool:
        s = line.strip()
        if not s.startswith('|'):
            return False
        return all(
            set(c.strip()) <= set('-: ')
            for c in s.strip('|').split('|')
            if c.strip()
        )

    def is_table_row(line: str) -> bool:
        s = line.strip()
        return s.startswith('|') and s.count('|') >= 2

    def is_label_cell(s: str) -> bool:
        """범주형 레이블로 쓸 수 있는 셀인지 판단 (숫자·수치 단위 포함 값 제외)."""
        if not s or s[0].isdigit():
            return False
        value_units = ('점', '원', '%', '학점', '이상', '이하', '회', '명', '개월', '년', '시간')
        return not any(s.endswith(u) for u in value_units)

    lines = text.split('\n')
    result = []
    i = 0

    while i < len(lines):
        line = lines[i]
        # 헤더 행 바로 다음이 구분자 행이면 Markdown 표로 처리
        if (is_table_row(line)
                and i + 1 < len(lines)
                and is_separator_row(lines[i + 1])):

            headers = parse_cells(line)
            result.append(line)          # 헤더 행 원본 유지
            result.append(lines[i + 1])  # 구분자 행 유지
            i += 2

            while i < len(lines) and is_table_row(lines[i]):
                cells = parse_cells(lines[i])
                while len(cells) < len(headers):
                    cells.append('')

                label = cells[0] if cells else ''
                use_label = bool(label) and is_label_cell(label)
                start = 1 if use_label else 0

                pairs = [
                    f"{h}: {v}"
                    for h, v in zip(headers[start:], cells[start:])
                    if v and h
                ]
                if pairs:
                    row_text = ' | '.join(pairs)
                    annotated = f"[{label}] {row_text}" if use_label else row_text
                    result.append(annotated)
                else:
                    result.append(lines[i])  # 변환 실패 시 원본 유지
                i += 1
        else:
            result.append(line)
            i += 1

    return '\n'.join(result)


def parse_pdf_batch(file_paths):
    """opendataloader_pdf로 여러 PDF를 한 번에 Markdown으로 변환.
    JVM 부팅이 호출당 ~수백ms라 배치가 필수. 반환: dict[절대경로]→markdown 텍스트.

    같은 basename을 가진 PDF가 다른 디렉터리에 있을 수 있으므로 (예: 서로 다른
    공지의 동명 첨부) 인덱스 prefix를 붙여 입력 디렉터리에 복사한 뒤 변환한다.
    """
    if not file_paths:
        return {}
    abs_paths = [os.path.abspath(p) for p in file_paths]
    results = {p: "" for p in abs_paths}
    with tempfile.TemporaryDirectory() as tmp:
        in_dir = os.path.join(tmp, "in")
        out_dir = os.path.join(tmp, "out")
        os.makedirs(in_dir)
        os.makedirs(out_dir)

        # 인덱스 prefix를 붙여 복사 → md 파일 stem으로 원본 경로를 역추적
        stem_to_orig = {}
        copied_inputs = []
        for i, p in enumerate(abs_paths):
            base = os.path.splitext(os.path.basename(p))[0]
            new_stem = f"{i:04d}_{base}"
            new_path = os.path.join(in_dir, new_stem + ".pdf")
            # 복사 실패는 환경 문제(권한/디스크 등)이므로 OSError 가 자연스럽게 전파됨
            shutil.copyfile(p, new_path)
            stem_to_orig[new_stem] = p
            copied_inputs.append(new_path)

        if not copied_inputs:
            return results

        try:
            # v1.0+: run(input_path=폴더, output_folder=..., generate_markdown=True)
            # 이전 버전 convert(input_path=리스트, ...) API는 제거됨
            opendataloader_pdf.run(
                input_path=in_dir,
                output_folder=out_dir,
                generate_markdown=True,
            )
        except (RuntimeError, OSError, ValueError, subprocess.CalledProcessError) as e:
            raise RuntimeError(
                f"opendataloader-pdf 변환 실패 ({len(copied_inputs)}개 PDF). "
                f"Java 11+ 설치 여부 확인 필요. 원본: {e}"
            ) from e

        for fname in os.listdir(out_dir):
            if not fname.endswith(".md"):
                continue
            stem = os.path.splitext(fname)[0]
            orig = stem_to_orig.get(stem)
            if not orig:
                continue
            try:
                with open(os.path.join(out_dir, fname), encoding="utf-8") as f:
                    text = f.read().strip()
                    text = _annotate_track_sections(text)
                    text = _annotate_markdown_table_rows(text)
                    results[orig] = text
            except OSError as e:
                print(f"  [PDF md 읽기 오류] {orig}: {e}")
    return results


def parse_pdf(file_path):
    """단일 PDF 파싱. 내부적으로 parse_pdf_batch 호출 (JVM 1회 부팅 비용 발생)."""
    abs_path = os.path.abspath(file_path)
    return parse_pdf_batch([abs_path]).get(abs_path, "")


def _libreoffice_convert(file_path, fmt, ext):
    """LibreOffice로 변환 후 결과 파일 경로 반환 (실패 시 None)"""
    if not LIBREOFFICE:
        print("  [오류] LibreOffice를 찾을 수 없습니다 (PATH 또는 표준 설치 경로 확인 필요).")
        return None
    try:
        subprocess.run([
            LIBREOFFICE,
            "--headless",
            "--convert-to", fmt,
            "--outdir", TMP_DIR,
            file_path
        ], timeout=60, capture_output=True)
        base_name = os.path.splitext(os.path.basename(file_path))[0]
        out_path = os.path.join(TMP_DIR, base_name + "." + ext)
        return out_path if os.path.exists(out_path) else None
    except subprocess.TimeoutExpired:
        print(f"  [HWP 타임아웃] {file_path}")
        return None
    except Exception as e:
        print(f"  [HWP 변환 오류] {file_path}: {e}")
        return None


def _hwp5txt_extract(file_path):
    """pyhwp의 hwp5txt로 HWP에서 텍스트 추출 (LibreOffice가 못 읽는 파일용 fallback)"""
    hwp5txt = shutil.which("hwp5txt")
    cmd = [hwp5txt, file_path] if hwp5txt else [
        sys.executable, "-m", "hwp5.hwp5txt", file_path
    ]
    try:
        res = subprocess.run(cmd, timeout=60, capture_output=True)
        if res.returncode != 0:
            return ""
        return res.stdout.decode("utf-8", errors="ignore").strip()
    except subprocess.TimeoutExpired:
        print(f"  [hwp5txt 타임아웃] {file_path}")
        return ""
    except Exception as e:
        print(f"  [hwp5txt 오류] {file_path}: {e}")
        return ""


def parse_hwp(file_path):
    """HWP/HWPX 파싱: LibreOffice→PDF→PyMuPDF 우선, 실패 시 pyhwp로 fallback.
    .hwpx는 pyhwp가 지원 안 하므로 LibreOffice 경로만 사용."""
    ext = os.path.splitext(file_path)[1].lower()

    # 1차: LibreOffice → PDF → PyMuPDF (표/레이아웃 보존에 유리)
    pdf_path = _libreoffice_convert(file_path, "pdf", "pdf")
    if pdf_path:
        try:
            text = parse_pdf(pdf_path)
        except Exception as e:
            print(f"  [HWP→PDF 추출 실패] {e}")
            text = ""
        finally:
            try:
                os.remove(pdf_path)
            except OSError:
                pass
        if text.strip():
            return text
        print(f"  [Fallback] HWP→PDF 결과 비어 있음 - 다음 방법 시도")

    # 2차: pyhwp (LibreOffice가 거부하는 HWP 파일에 강함). .hwpx는 미지원.
    if ext == ".hwp":
        text = _hwp5txt_extract(file_path)
        if text:
            print(f"  [pyhwp로 추출 성공]")
            return text

    # 3차: LibreOffice → TXT
    txt_path = _libreoffice_convert(file_path, "txt:Text", "txt")
    if not txt_path:
        return ""
    try:
        with open(txt_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
        return text.strip()
    except Exception as e:
        print(f"  [TXT 읽기 실패] {e}")
        return ""
    finally:
        try:
            os.remove(txt_path)
        except OSError:
            pass


def parse_xlsx(file_path):
    """엑셀 파일에서 텍스트 추출. 각 셀에 컬럼 헤더를 붙여 청킹 후에도 의미 보존."""
    result = []
    try:
        xl = pd.ExcelFile(file_path)
        for sheet_name in xl.sheet_names:
            df = xl.parse(sheet_name).fillna("")
            result.append(f"[시트: {sheet_name}]")
            for _, row in df.iterrows():
                cells = []
                for col, val in row.items():
                    col_str = str(col).strip()
                    val_str = str(val).strip()
                    if not val_str:
                        continue
                    # pandas unnamed 컬럼('Unnamed: N')은 헤더 없는 셀 — 값만 출력
                    if col_str and not col_str.startswith('Unnamed:'):
                        cells.append(f"{col_str}: {val_str}")
                    else:
                        cells.append(val_str)
                if cells:
                    result.append(" | ".join(cells))
    except Exception as e:
        print(f"  [XLSX 오류] {file_path}: {e}")
    return "\n".join(result)


def parse_docx(file_path):
    """DOCX에서 텍스트 추출"""
    result = []
    try:
        doc = Document(file_path)
        for para in doc.paragraphs:
            if para.text.strip():
                result.append(para.text.strip())
        for table in doc.tables:
            header_cells: list = []
            for row_idx, row in enumerate(table.rows):
                seen_cells: set = set()
                cells: list = []
                for cell in row.cells:
                    txt = cell.text.strip()
                    if txt and txt not in seen_cells:
                        cells.append(txt)
                        seen_cells.add(txt)
                if not cells:
                    continue
                if row_idx == 0:
                    header_cells = cells
                    result.append("[표헤더] " + " | ".join(cells))
                else:
                    if header_cells and len(header_cells) == len(cells):
                        parts = [f"{h}: {v}" for h, v in zip(header_cells, cells)]
                    else:
                        parts = cells
                    result.append("[표] " + " | ".join(parts))
    except Exception as e:
        print(f"  [DOCX 오류] {file_path}: {e}")
    return "\n".join(result)


def parse_zip(file_path):
    """ZIP 압축 해제 후 내부 파일 파싱. 내부 PDF는 한 번에 배치 처리."""
    result = []
    tmp_zip_dir = os.path.join(BASE_DIR, "data", "tmp_zip")
    try:
        os.makedirs(tmp_zip_dir, exist_ok=True)
        # zip-slip 방지: 추출 경로가 base 디렉터리를 벗어나는 엔트리는 거부
        base_dir = os.path.realpath(tmp_zip_dir) + os.sep
        with zipfile.ZipFile(file_path, "r") as z:
            for member in z.infolist():
                target = os.path.realpath(os.path.join(tmp_zip_dir, member.filename))
                if not (target + os.sep).startswith(base_dir):
                    print(f"  [ZIP 위험 엔트리 거부] {member.filename}")
                    continue
                z.extract(member, tmp_zip_dir)

        # 내부 파일 수집 (확장자별 분류)
        inner_files = []
        for root, _, files in os.walk(tmp_zip_dir):
            for fname in files:
                inner_files.append((fname, os.path.join(root, fname)))
        pdf_paths = [p for f, p in inner_files if os.path.splitext(f)[1].lower() == ".pdf"]
        pdf_results = parse_pdf_batch(pdf_paths) if pdf_paths else {}

        for fname, inner_path in inner_files:
            ext = os.path.splitext(fname)[1].lower()
            text = ""
            if ext == ".pdf":
                text = pdf_results.get(os.path.abspath(inner_path), "")
            elif ext in [".hwp", ".hwpx"]:
                text = parse_hwp(inner_path)
            elif ext in [".xlsx", ".xls"]:
                text = parse_xlsx(inner_path)
            elif ext == ".docx":
                text = parse_docx(inner_path)

            if text:
                # ZIP 안에 같은 이름이 다른 디렉터리에 있을 수 있어 상대 경로 사용
                rel = os.path.relpath(inner_path, tmp_zip_dir)
                result.append(f"[ZIP 내부: {rel}]\n{text}")

    except RuntimeError:
        # parse_pdf_batch가 raise한 환경 오류는 상위로 그대로 전파
        raise
    except Exception as e:
        print(f"  [ZIP 오류] {file_path}: {e}")
    finally:
        shutil.rmtree(tmp_zip_dir, ignore_errors=True)

    return "\n\n".join(result)


def parse_file(file_path):
    """파일 확장자에 따라 적절한 파서 호출"""
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return parse_pdf(file_path)
    elif ext in [".hwp", ".hwpx"]:
        return parse_hwp(file_path)
    elif ext in [".xlsx", ".xls"]:
        return parse_xlsx(file_path)
    elif ext == ".docx":
        return parse_docx(file_path)
    elif ext == ".zip":
        return parse_zip(file_path)
    else:
        print(f"  [스킵] 지원하지 않는 형식: {file_path}")
        return ""


def warn_if_empty(file_name, text):
    """추출 텍스트가 너무 적으면 경고 (스캔본 PDF 등 의심)"""
    if not text:
        print(f"  [!경고] 추출 실패 (0자): {file_name}")
    elif len(text) < EMPTY_TEXT_THRESHOLD:
        print(f"  [!경고] 추출 텍스트 매우 적음 ({len(text)}자) - 스캔본 가능성: {file_name}")


def parse_all():
    """크롤링 첨부파일 파싱 (증분: 기존 결과가 있으면 재사용)"""
    notices_path = os.path.join(RAW_DIR, "notices.json")
    output_path = os.path.join(PARSED_DIR, "notices_parsed.json")

    with open(notices_path, encoding="utf-8") as f:
        notices = json.load(f)

    # 기존 파싱 결과 로드 (증분 파싱)
    existing = {}
    if os.path.exists(output_path):
        try:
            with open(output_path, encoding="utf-8") as f:
                prev = json.load(f)
            for n in prev:
                try:
                    wr = n["url"].split("wr_id=")[1].split("&")[0]
                except (KeyError, IndexError):
                    continue
                existing[wr] = {
                    a["name"]: a.get("parsed_text", "")
                    for a in n.get("attachments", []) if a.get("name")
                }
        except Exception:
            existing = {}

    print(f"[시작] 첨부파일 파싱 (공지 {len(notices)}개)")
    cached_count = 0
    parsed_count = 0

    # PDF 배치 변환을 위해 처리 대상 PDF 경로 먼저 수집
    pdf_to_process = []
    for notice in notices:
        try:
            wr_id = notice["url"].split("wr_id=")[1].split("&")[0]
        except (KeyError, IndexError):
            continue
        attach_dir = os.path.join(ATTACHMENT_DIR, wr_id)
        for att in notice.get("attachments", []):
            file_name = att["name"]
            ext = os.path.splitext(file_name)[1].lower()
            if ext != ".pdf":
                continue
            cached = existing.get(wr_id, {}).get(file_name, "")
            if cached:
                continue
            file_path = os.path.join(attach_dir, file_name)
            if os.path.exists(file_path):
                pdf_to_process.append(file_path)

    pdf_results = {}
    if pdf_to_process:
        print(f"[PDF 배치] {len(pdf_to_process)}개 PDF를 opendataloader로 변환 중...")
        pdf_results = parse_pdf_batch(pdf_to_process)

    for notice in notices:
        try:
            wr_id = notice["url"].split("wr_id=")[1].split("&")[0]
        except (KeyError, IndexError):
            continue
        attach_dir = os.path.join(ATTACHMENT_DIR, wr_id)
        attachments = notice.get("attachments", [])

        if not attachments:
            continue

        print(f"\n[{wr_id}] {notice['title'][:40]}")

        for att in attachments:
            file_name = att["name"]
            file_path = os.path.join(attach_dir, file_name)
            ext = os.path.splitext(file_name)[1].lower()

            cached = existing.get(wr_id, {}).get(file_name, "")
            if cached:
                att["parsed_text"] = cached
                cached_count += 1
                print(f"  [캐시] {file_name} ({len(cached)}자)")
                continue

            if not os.path.exists(file_path):
                print(f"  [없음] {file_name}")
                continue

            if ext == ".pdf":
                text = pdf_results.get(os.path.abspath(file_path), "")
            else:
                print(f"  파싱 중: {file_name}")
                text = parse_file(file_path)
            att["parsed_text"] = text
            warn_if_empty(file_name, text)
            print(f"  완료: {len(text)}자 추출")
            parsed_count += 1

        # 공지 단위로 저장 (크래시 시 이전 공지 결과 보존)
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(notices, f, ensure_ascii=False, indent=2)

    total_files = sum(len(n.get("attachments", [])) for n in notices)
    success = sum(
        1 for n in notices for a in n.get("attachments", []) if a.get("parsed_text")
    )

    print(f"\n[완료] 파싱 결과 저장 → {output_path}")
    print(f"신규 파싱: {parsed_count}개 / 캐시 재사용: {cached_count}개 / 성공: {success}/{total_files}개")
    _cleanup_tmp()


def parse_manual_files():
    """manual_files 폴더의 파일들 파싱 (증분)"""
    manual_dir = os.path.join(BASE_DIR, "data", "manual_files")
    output_path = os.path.join(PARSED_DIR, "manual_parsed.json")

    existing = {}
    if os.path.exists(output_path):
        try:
            with open(output_path, encoding="utf-8") as f:
                prev = json.load(f)
            existing = {r["file_name"]: r.get("parsed_text", "") for r in prev if r.get("file_name")}
        except Exception:
            existing = {}

    results = []
    cached_count = 0
    parsed_count = 0
    print(f"[시작] manual_files 파싱")

    # PDF 배치 변환 대상 수집
    pdf_to_process = []
    for fname in sorted(os.listdir(manual_dir)):
        ext = os.path.splitext(fname)[1].lower()
        if ext == ".pdf" and not existing.get(fname):
            pdf_to_process.append(os.path.join(manual_dir, fname))

    pdf_results = {}
    if pdf_to_process:
        print(f"[PDF 배치] {len(pdf_to_process)}개 PDF를 opendataloader로 변환 중...")
        pdf_results = parse_pdf_batch(pdf_to_process)

    for fname in sorted(os.listdir(manual_dir)):
        fpath = os.path.join(manual_dir, fname)
        ext = os.path.splitext(fname)[1].lower()

        if ext not in [".pdf", ".hwp", ".hwpx", ".xlsx", ".xls", ".docx", ".txt", ".zip"]:
            print(f"  [스킵] {fname}")
            continue

        cached_text = existing.get(fname)
        if cached_text:
            print(f"  [캐시] {fname} ({len(cached_text)}자)")
            results.append({"file_name": fname, "parsed_text": cached_text})
            cached_count += 1
            continue

        if ext == ".pdf":
            text = pdf_results.get(os.path.abspath(fpath), "")
        elif ext == ".txt":
            print(f"  파싱 중: {fname}")
            with open(fpath, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read().strip()
        else:
            print(f"  파싱 중: {fname}")
            text = parse_file(fpath)

        warn_if_empty(fname, text)
        print(f"  완료: {len(text)}자 추출")
        results.append({"file_name": fname, "parsed_text": text})
        parsed_count += 1

    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(results, f, ensure_ascii=False, indent=2)

    success = sum(1 for r in results if r["parsed_text"])
    print(f"\n[완료] 저장 → {output_path}")
    print(f"신규 파싱: {parsed_count}개 / 캐시 재사용: {cached_count}개 / 성공: {success}/{len(results)}개")
    _cleanup_tmp()


def _cleanup_tmp():
    """LibreOffice 임시 산출물 정리"""
    if os.path.exists(TMP_DIR):
        for f in os.listdir(TMP_DIR):
            try:
                os.remove(os.path.join(TMP_DIR, f))
            except OSError:
                pass


if __name__ == "__main__":
    if not LIBREOFFICE:
        print("[경고] LibreOffice를 찾지 못했습니다. HWP/HWPX 파싱이 실패할 수 있습니다.")
        print("       - macOS/Linux: 'soffice' 명령이 PATH에 있는지 확인")
        print("       - Windows: C:\\Program Files\\LibreOffice\\program\\soffice.exe 설치 확인")
    else:
        print(f"[LibreOffice] {LIBREOFFICE}")

    # 1. 기존 공지사항 첨부파일 파싱
    # parse_all()

    # 2. 수동 업로드 파일 파싱
    parse_manual_files()
